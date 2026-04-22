import os
import sys
import json
import zipfile
import re
import xml.etree.ElementTree as ET

from openai import OpenAI
from docx import Document
from docx.shared import Pt
from config_manager import get_api_key

CONFIG_FILENAME = "config.json"


def app_dir():
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


def config_path():
    return os.path.join(app_dir(), CONFIG_FILENAME)


def get_current_project_dir():
    path = config_path()
    if not os.path.exists(path):
        return ""

    try:
        with open(path, "r", encoding="utf-8") as f:
            config = json.load(f)
            return config.get("current_project_dir", "").strip()
    except Exception:
        return ""


def find_transcript_file(project_dir):
    txt_file = None
    docx_file = None

    for f in os.listdir(project_dir):
        lower_name = f.lower()
        full_path = os.path.join(project_dir, f)

        if lower_name.endswith(".txt"):
            txt_file = full_path
        elif lower_name.endswith(".docx"):
            docx_file = full_path

    if txt_file:
        return txt_file
    if docx_file:
        return docx_file

    return None


def read_txt_file(path):
    encodings = ["utf-8", "utf-8-sig", "gbk", "gb18030"]

    for enc in encodings:
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except Exception:
            continue

    raise RuntimeError(f"Could not read txt file encoding: {path}")


def read_docx_file(path):
    try:
        with zipfile.ZipFile(path, "r") as z:
            xml_content = z.read("word/document.xml")
        root = ET.fromstring(xml_content)

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        texts = []

        for para in root.findall(".//w:p", ns):
            line_parts = []
            for node in para.findall(".//w:t", ns):
                if node.text:
                    line_parts.append(node.text)
            if line_parts:
                texts.append("".join(line_parts))

        return "\n".join(texts)
    except Exception:
        raise RuntimeError(
            "Failed to read docx: this file may not be a valid Word .docx file.\n"
            "Save the transcript as a standard .docx file, or use a .txt file instead."
        )


def load_transcript_content(transcript_path):
    lower_path = transcript_path.lower()

    if lower_path.endswith(".txt"):
        return read_txt_file(transcript_path)

    if lower_path.endswith(".docx"):
        return read_docx_file(transcript_path)

    raise RuntimeError(f"Unsupported transcript format: {transcript_path}")


def clean_transcript_text(text):
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    return "\n".join(lines)


def build_prompt(transcript_text, transcript_name, has_pdf):
    pdf_hint = "slides.pdf was found" if has_pdf else "slides.pdf was not found"

    return f"""
Please turn the following course transcript into clear, structured English study notes that are easy to review.

Requirements:
1. Output in English.
2. Format the content so it works well as Word study notes.
3. Do not use Markdown tables.
4. Do not use HTML tags such as <br>.
5. Use clear section headings.
6. Include these sections when possible:
   Course Topic
   Core Concepts
   Key Points
   Important Examples/Cases (if any)
   Summary
7. Remove obvious repetition, filler, and meaningless speech fragments.
8. Preserve genuinely useful information and make the notes read like formal study notes.
9. If the transcript is unclear or recognition is messy, summarize based on context instead of copying garbled text.
10. Do not invent information that is not in the transcript; if information is limited, organize the available content as well as possible.

Additional information:
- Transcript file: {transcript_name}
- PPT PDF status: {pdf_hint}

Transcript:
{transcript_text}
""".strip()


def clean_summary_output(text):
    replacements = {
        "<br>": "\n",
        "<br/>": "\n",
        "<br />": "\n",
        "### ": "",
        "## ": "",
        "# ": "",
        "**": "",
        "✓": "•",
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    # Remove markdown table separators
    lines = text.splitlines()
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if re.fullmatch(r"\|?[-:\s|]+\|?", stripped):
            continue
        cleaned.append(line)

    # Collapse excessive blank lines
    final_lines = []
    previous_blank = False
    for line in cleaned:
        is_blank = not line.strip()
        if is_blank:
            if not previous_blank:
                final_lines.append("")
            previous_blank = True
        else:
            final_lines.append(line.rstrip())
            previous_blank = False

    return "\n".join(final_lines).strip()


def is_main_heading(line):
    candidates = {
        "Course Topic",
        "Core Concepts",
        "Key Points",
        "Important Examples",
        "Important Examples/Cases",
        "Cases",
        "Summary",
    }
    pure = line.strip().strip("：:").strip()
    return pure in candidates


def save_to_word(summary_text, output_path):
    document = Document()

    # Base font
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    document.add_heading("Course Study Notes", level=0)

    for raw_line in summary_text.splitlines():
        line = raw_line.strip()

        if not line:
            document.add_paragraph("")
            continue

        # Top-level heading
        if is_main_heading(line):
            document.add_heading(line.strip("：:"), level=1)
            continue

        # Lines shaped like "Course Topic: xxx"
        if "：" in line:
            left, right = line.split("：", 1)
            if is_main_heading(left):
                document.add_heading(left.strip(), level=1)
                if right.strip():
                    document.add_paragraph(right.strip())
                continue

        if ":" in line:
            left, right = line.split(":", 1)
            if is_main_heading(left):
                document.add_heading(left.strip(), level=1)
                if right.strip():
                    document.add_paragraph(right.strip())
                continue

        # List item
        if line.startswith(("•", "-", "1.", "2.", "3.", "4.", "5.")):
            document.add_paragraph(line, style=None)
            continue

        document.add_paragraph(line)

    document.save(output_path)


def main():
    api_key = get_api_key()
    if not api_key:
        print("API Key was not found. Set the API Key in the app first.")
        sys.exit(1)

    project_dir = get_current_project_dir()
    if not project_dir:
        print("Current project path was not found. Click \"Start Recording\" in the main app first to create a project.")
        sys.exit(1)

    pdf_path = os.path.join(project_dir, "slides.pdf")
    transcript_path = find_transcript_file(project_dir)
    output_path = os.path.join(project_dir, "summary.docx")

    if not os.path.exists(pdf_path):
        print("slides.pdf was not found. Generate the PDF first.")
        sys.exit(1)

    if not transcript_path:
        print("No transcript was found (txt or docx). Place it in the project folder.")
        sys.exit(1)

    try:
        transcript_content = load_transcript_content(transcript_path)
    except Exception as e:
        print(f"Failed to read transcript: {e}")
        sys.exit(1)

    transcript_content = clean_transcript_text(transcript_content)

    if not transcript_content.strip():
        print("The transcript is empty, so it cannot be summarized.")
        sys.exit(1)

    transcript_content = transcript_content[:20000]

    prompt = build_prompt(
        transcript_text=transcript_content,
        transcript_name=os.path.basename(transcript_path),
        has_pdf=os.path.exists(pdf_path)
    )

    try:
        client = OpenAI(
            api_key=api_key,
            base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
        )

        resp = client.chat.completions.create(
            model="qwen-plus",
            messages=[
                {"role": "system", "content": "You are an assistant skilled at organizing course content and creating study notes."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )

        summary = resp.choices[0].message.content.strip()
        summary = clean_summary_output(summary)

        save_to_word(summary, output_path)

        print(f"Summary complete. Saved to: {output_path}")
        sys.exit(0)

    except Exception as e:
        print(f"Summary failed: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
